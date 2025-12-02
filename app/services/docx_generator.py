from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re


# -------------------------------------------------------
# UTILITY: ADD HYPERLINK
# -------------------------------------------------------

def add_hyperlink(paragraph, url, text, font_size=10):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    run.append(rPr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    run.append(text_elem)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)

    return hyperlink


# -------------------------------------------------------
# NEW DOCUMENT
# -------------------------------------------------------

def new_document():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)

    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    return doc


# -------------------------------------------------------
# FORCE SPACING = 0
# -------------------------------------------------------

def force_zero_spacing(p):
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


# -------------------------------------------------------
# HEADER BLOCK
# -------------------------------------------------------

def add_resume_header(doc, user_info):
    # NAME
    p = doc.add_paragraph()
    run = p.add_run(user_info.get("full_name", "").upper())
    run.bold = True
    run.font.size = Pt(16)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    force_zero_spacing(p)

    # CONTACT LINE
    p2 = doc.add_paragraph()
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.space_before = Pt(0)
    force_zero_spacing(p2)

    parts_added = False

    if user_info.get("phone"):
        p2.add_run(user_info["phone"])
        parts_added = True

    if user_info.get("email"):
        if parts_added: p2.add_run(" | ")
        p2.add_run(user_info["email"])
        parts_added = True

    if user_info.get("linkedin"):
        if parts_added: p2.add_run(" | ")
        add_hyperlink(p2, user_info["linkedin"], "LinkedIn")
        parts_added = True

    if user_info.get("github"):
        if parts_added: p2.add_run(" | ")
        add_hyperlink(p2, user_info["github"], "GitHub")

    # DIVIDER
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(0)
    p3.paragraph_format.space_before = Pt(0)
    force_zero_spacing(p3)

    p3._p.get_or_add_pPr()
    pPr = p3._p.pPr

    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), "000000")
    pbdr.append(bottom)

    pPr.append(pbdr)


# -------------------------------------------------------
# SECTION HEADER
# -------------------------------------------------------

def add_section_header(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)

    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1

    force_zero_spacing(p)

    p._p.get_or_add_pPr()
    pPr = p._p.pPr

    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), "000000")
    pbdr.append(bottom)
    pPr.append(pbdr)

    return p


# -------------------------------------------------------
# HELPERS
# -------------------------------------------------------

def normalize_bullet(text):
    if not text:
        return "• "
    clean = re.sub(r'^[\s•\-\*\–\—\·\●\▪\○\▶\›\t]+', "", text).strip()
    return f"• {clean}"


# -------------------------------------------------------
# SECTIONS
# -------------------------------------------------------

def add_summary(doc, summary):
    p = doc.add_paragraph(summary)
    p.paragraph_format.space_after = Pt(10)
    return doc


def add_experience(doc, experience_list):
    section = doc.sections[0]
    text_width = section.page_width - section.left_margin - section.right_margin

    for job in experience_list:
        p = doc.add_paragraph()
        p_format = p.paragraph_format
        p_format.space_after = Pt(0)

        # Title + Company
        p.add_run(f"{job.get('title', '')}, {job.get('company', '')}").bold = True

        # Dates aligned to the SAME RIGHT COLUMN
        tab_pos = text_width - Inches(1.7)
        p_format.tab_stops.add_tab_stop(tab_pos)

        p.add_run("\t")
        p.add_run(job.get("dates", "")).bold = True

        last_bullet = None
        for bullet in job.get("bullets", []):
            clean = normalize_bullet(bullet)
            bp = doc.add_paragraph(clean)
            bp_format = bp.paragraph_format
            bp_format.left_indent = Pt(12)
            bp_format.first_line_indent = Pt(-6)
            bp_format.space_after = Pt(0)
            last_bullet = bp

        if last_bullet:
            last_bullet.paragraph_format.space_after = Pt(6)
        else:
            p_format.space_after = Pt(6)

    return doc


def add_education(doc, education_list):
    section = doc.sections[0]
    text_width = section.page_width - section.left_margin - section.right_margin

    for edu in education_list:
        p = doc.add_paragraph()
        p_format = p.paragraph_format

        # Degree + Institution (LEFT SIDE)
        left_text = f"{edu.get('degree', '')}, {edu.get('institution', '')}"
        p.add_run(left_text)

        # Right-aligned dates (MATCH EXPERIENCE)
        tab_pos = text_width - Inches(1.7)
        p_format.tab_stops.add_tab_stop(tab_pos)

        p.add_run("\t")

        date_text = edu.get("dates", "")
        p.add_run(date_text).bold = True

        p_format.space_after = Pt(10)

    return doc


def add_skills(doc, skills_dict):
    for category, skills in skills_dict.items():

        p = doc.add_paragraph()
        r1 = p.add_run(f"{category}: ")
        r1.bold = True

        r2 = p.add_run(", ".join(skills))
        p.paragraph_format.space_after = Pt(0)

    return doc


# -------------------------------------------------------
# MAIN GENERATOR
# -------------------------------------------------------

def generate_final_docx(sections, user_info, output_path):
    doc = new_document()

    add_resume_header(doc, user_info)

    add_section_header(doc, "SUMMARY")
    add_summary(doc, sections["summary"])

    add_section_header(doc, "EXPERIENCE")
    add_experience(doc, sections["experience"])

    add_section_header(doc, "EDUCATION")
    add_education(doc, sections["education"])

    add_section_header(doc, "SKILLS")
    add_skills(doc, sections["skills"])

    doc.save(output_path)
    return output_path